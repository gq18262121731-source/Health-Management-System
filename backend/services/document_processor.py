"""文档处理服务 - 支持多种格式的文档解析"""
import logging
from typing import Optional, Dict, Any
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，支持多种格式"""
    
    @staticmethod
    def process_file(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        处理文件并提取文本内容
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，自动检测）
        
        Returns:
            包含title和content的字典
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 自动检测文件类型
        if file_type is None:
            file_type = path.suffix.lower()
        
        # 根据文件类型选择处理方式
        if file_type in ['.txt', '.text']:
            return DocumentProcessor._process_text(path)
        elif file_type in ['.md', '.markdown']:
            return DocumentProcessor._process_markdown(path)
        elif file_type == '.pdf':
            return DocumentProcessor._process_pdf(path)
        elif file_type in ['.docx', '.doc']:
            return DocumentProcessor._process_docx(path)
        else:
            # 默认按文本处理
            logger.warning(f"未知文件类型 {file_type}，按文本处理")
            return DocumentProcessor._process_text(path)
    
    @staticmethod
    def _process_text(path: Path) -> Dict[str, Any]:
        """处理文本文件"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 使用文件名作为标题
            title = path.stem
            
            return {
                "title": title,
                "content": content
            }
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(path, 'r', encoding='gbk') as f:
                content = f.read()
            return {
                "title": path.stem,
                "content": content
            }
    
    @staticmethod
    def _process_markdown(path: Path) -> Dict[str, Any]:
        """处理Markdown文件"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 尝试从Markdown中提取标题
            title = path.stem
            lines = content.split('\n')
            for line in lines[:10]:  # 只检查前10行
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
                elif line.startswith('## '):
                    title = line[3:].strip()
                    break
            
            return {
                "title": title,
                "content": content
            }
        except Exception as e:
            logger.error(f"处理Markdown失败: {str(e)}")
            return DocumentProcessor._process_text(path)
    
    @staticmethod
    def _process_pdf(path: Path) -> Dict[str, Any]:
        """处理PDF文件"""
        try:
            import PyPDF2
            
            with open(path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # 提取所有文本
                content_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
                
                content = '\n\n'.join(content_parts)
                title = path.stem
                
                return {
                    "title": title,
                    "content": content
                }
        except ImportError:
            logger.error("PyPDF2未安装，无法处理PDF文件")
            raise ImportError("请安装PyPDF2: pip install PyPDF2")
        except Exception as e:
            logger.error(f"处理PDF失败: {str(e)}")
            raise
    
    @staticmethod
    def _process_docx(path: Path) -> Dict[str, Any]:
        """处理Word文档"""
        try:
            from docx import Document
            
            doc = Document(path)
            
            # 提取所有段落文本
            content_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            content = '\n\n'.join(content_parts)
            title = path.stem
            
            # 尝试使用第一个段落作为标题
            if doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if first_para and len(first_para) < 100:
                    title = first_para
            
            return {
                "title": title,
                "content": content
            }
        except ImportError:
            logger.error("python-docx未安装，无法处理Word文档")
            raise ImportError("请安装python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"处理Word文档失败: {str(e)}")
            raise


